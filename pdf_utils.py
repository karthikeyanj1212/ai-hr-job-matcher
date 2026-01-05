"""
PDF and DOCX text extraction utilities
Fixed for Streamlit file upload handling
"""

import re
import io


def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file."""
    text = ""
    
    # Reset file pointer
    try:
        file.seek(0)
    except:
        pass
    
    # Read file content into bytes
    try:
        file_bytes = file.read()
        file.seek(0)  # Reset for potential reuse
    except Exception as e:
        return f"Error reading file: {e}"
    
    # Try pdfplumber first (better for complex PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return clean_text(text)
    except ImportError:
        pass
    except Exception as e:
        print(f"pdfplumber error: {e}")
    
    # Fallback to PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        if text.strip():
            return clean_text(text)
    except ImportError:
        pass
    except Exception as e:
        print(f"PyPDF2 error: {e}")
    
    # Fallback to pymupdf (fitz)
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        if text.strip():
            return clean_text(text)
    except ImportError:
        pass
    except Exception as e:
        print(f"pymupdf error: {e}")
    
    return text or "Could not extract text from PDF. Please install: pip install pdfplumber PyPDF2 pymupdf"


def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file."""
    try:
        file.seek(0)
    except:
        pass
    
    try:
        from docx import Document
        
        # Read into BytesIO
        file_bytes = file.read()
        file.seek(0)
        
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text += " | ".join(row_text) + "\n"
        
        return clean_text(text)
    except ImportError:
        return "Please install python-docx: pip install python-docx"
    except Exception as e:
        return f"Error extracting DOCX: {e}"


def extract_text_from_txt(file) -> str:
    """Extract text from TXT file."""
    try:
        file.seek(0)
    except:
        pass
    
    try:
        content = file.read()
        if isinstance(content, bytes):
            try:
                return content.decode('utf-8')
            except:
                return content.decode('latin-1')
        return str(content)
    except Exception as e:
        return f"Error reading TXT file: {e}"


def clean_text(text: str) -> str:
    """Clean extracted text."""
    if not text:
        return ""
    
    # Remove excessive newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    # Remove excessive spaces
    text = re.sub(r' +', ' ', text)
    # Remove null characters
    text = text.replace('\x00', '')
    # Remove form feeds
    text = text.replace('\x0c', '\n')
    
    return text.strip()


def extract_text_from_file(file, filename: str) -> str:
    """Extract text based on file type."""
    if file is None:
        return ""
    
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.pdf'):
            return extract_text_from_pdf(file)
        elif filename_lower.endswith('.docx'):
            return extract_text_from_docx(file)
        elif filename_lower.endswith('.txt'):
            return extract_text_from_txt(file)
        else:
            return "Unsupported format. Please upload PDF, DOCX, or TXT."
    except Exception as e:
        return f"Error processing file: {e}"